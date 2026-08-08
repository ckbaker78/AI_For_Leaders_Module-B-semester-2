from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "CalvinBaker_DX699O2_Final_Project_Module_B.docx"

MORTGAGE_HEATMAP = (
    ROOT / "tmp/notebooks/extracted/Week9CalvinBakerSubmission/figure-004.png"
)
LENDING_HEATMAP = (
    ROOT / "tmp/notebooks/extracted/Week9CalvinBakerSubmission/figure-008.png"
)
MORTGAGE_BUBBLE = (
    ROOT / "tmp/notebooks/extracted/Week9CalvinBakerSubmission/figure-005.png"
)
PCA_FIGURE = ROOT / "tmp/report/figures/figure-04-pca-comparison.png"
IMPORTANCE_FIGURE = ROOT / "tmp/report/figures/figure-05-feature-importance.png"
DECILE_FIGURE = ROOT / "tmp/report/figures/figure-06-risk-deciles.png"

for figure in (
    MORTGAGE_HEATMAP,
    LENDING_HEATMAP,
    MORTGAGE_BUBBLE,
    PCA_FIGURE,
    IMPORTANCE_FIGURE,
    DECILE_FIGURE,
):
    if not figure.exists():
        raise FileNotFoundError(figure)


def set_run_font(run, *, size=12, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def format_paragraph(paragraph, *, align=None, double=True, after=0, before=0):
    fmt = paragraph.paragraph_format
    if double:
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        fmt.line_spacing = 2
    else:
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.line_spacing = 1
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.widow_control = True
    if align is not None:
        paragraph.alignment = align


def add_text(text="", *, style=None, align=None, bold=False, italic=False):
    paragraph = doc.add_paragraph(style=style)
    format_paragraph(paragraph, align=align)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return paragraph


def add_rich(segments, *, style=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    format_paragraph(paragraph, align=align)
    for text, options in segments:
        run = paragraph.add_run(text)
        set_run_font(
            run,
            bold=options.get("bold"),
            italic=options.get("italic"),
            color=options.get("color"),
        )
    return paragraph


def add_heading(text, *, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    format_paragraph(paragraph)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, bold=True, italic=(level == 2))
    return paragraph


def add_bullet(text, *, level=0):
    paragraph = doc.add_paragraph()
    format_paragraph(paragraph)
    paragraph.style = doc.styles["List Bullet"]
    paragraph.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_caption(text):
    paragraph = doc.add_paragraph(style="Figure Caption")
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.keep_with_next = False
    run = paragraph.add_run(text)
    set_run_font(run, italic=True)
    return paragraph


def add_picture(path, *, width):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    return paragraph


def add_two_panel_figures(left_path, left_caption, right_path, right_caption):
    figure_table = doc.add_table(rows=1, cols=2)
    figure_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    figure_table.autofit = False
    for cell, path, caption in (
        (figure_table.rows[0].cells[0], left_path, left_caption),
        (figure_table.rows[0].cells[1], right_path, right_caption),
    ):
        cell.width = Inches(3.25)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cell.text = ""
        picture_paragraph = cell.paragraphs[0]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.space_before = Pt(0)
        picture_paragraph.paragraph_format.space_after = Pt(0)
        picture_paragraph.add_run().add_picture(str(path), width=Inches(3.08))
        caption_paragraph = cell.add_paragraph()
        format_paragraph(caption_paragraph, align=WD_ALIGN_PARAGRAPH.CENTER)
        caption_run = caption_paragraph.add_run(caption)
        set_run_font(caption_run, italic=True)

    table_pr = figure_table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "nil")
        borders.append(element)
    table_pr.append(borders)
    return figure_table


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, align=align)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_table_borders(table, color="B7BDC7", size="6"):
    table_pr = table._tbl.tblPr
    borders = table_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    r_pr.append(r_fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F4E79")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_run_font(run)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, fallback, end])


def page_break():
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
normal.paragraph_format.space_after = Pt(0)

for name in ("Heading 1", "Heading 2"):
    style = styles[name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.keep_with_next = True

if "Figure Caption" not in styles:
    caption_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
else:
    caption_style = styles["Figure Caption"]
caption_style.font.name = "Times New Roman"
caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
caption_style.font.size = Pt(12)
caption_style.font.italic = True
caption_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

list_bullet = styles["List Bullet"]
list_bullet.font.name = "Times New Roman"
list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
list_bullet.font.size = Pt(12)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
format_paragraph(footer, align=WD_ALIGN_PARAGRAPH.CENTER, double=False)
footer_run = footer.add_run("Calvin Baker | DX699O2 Final Project | Page ")
set_run_font(footer_run, size=10)
add_page_number(footer)

doc.core_properties.title = (
    "Credit-Risk Segmentation and Predictive Modeling: Module B Capstone Proposal"
)
doc.core_properties.subject = "DX699O2 Final Project, Summer 2026"
doc.core_properties.author = "Calvin Baker"

# ---------------------------------------------------------------------------
# BODY PAGE 1 — Title and required Section 1
# ---------------------------------------------------------------------------
add_text(
    "Credit-Risk Segmentation and Predictive Modeling: Module B Capstone Proposal",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    bold=True,
)
add_text(
    "Calvin Baker | DX699O2 | Summer 2026",
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_heading("1. Problem statement/description")
add_text(
    "Consumer lenders need to identify loans that merit additional review before losses "
    "materialize, while avoiding the opposite error of treating creditworthy borrowers as "
    "high risk. This project asks whether routinely available application and loan features "
    "can rank adverse outcomes well enough to support a transparent, capacity-limited human "
    "review process. The business value is earlier monitoring and better allocation of "
    "underwriting attention; the public value is a more consistent process with explicit "
    "limits on automated decision-making."
)
add_text(
    "The analysis compares two distinct populations. The mortgage Loan Default dataset has "
    "148,670 records and a binary Status outcome with a 24.6% overall adverse rate. The "
    "Accepted Lending Club file contains 2,260,701 issued-loan records; the modeling sample "
    "uses 61,100 evenly spaced rows and defines a bad loan from completed adverse statuses, "
    "yielding a 13.0% adverse rate. Because the targets are related but not identical, "
    "performance is interpreted within each dataset rather than as a lender-to-lender scorecard."
)
add_text(
    "Weeks 1–7 established data quality, skew, class imbalance, and pairwise relationships. "
    "Weeks 8–12 added interactions, PCA, logistic baselines, and tuned random forests. The "
    "models offer useful ranking but not automatic approval or denial: mortgage missingness "
    "is outcome-linked, Lending Club includes only accepted applicants, and neither result "
    "is an out-of-time production validation."
)

# ---------------------------------------------------------------------------
# BODY PAGE 2 — Required Section 2a, heatmaps
# ---------------------------------------------------------------------------
page_break()
add_heading("2. Exploratory data analysis")
add_heading("a. Multivariate analysis", level=2)
add_text(
    "Binned heatmaps expose interactions without imposing a linear trend. Mortgage default "
    "is elevated in the 100%+ loan-to-value (LTV) tier across debt-to-income (DTI) bands, "
    "although DTI is nonlinear and edge cells are smaller. In Lending Club, bad-loan rates "
    "generally rise from grade A through G across purposes; purpose is secondary and sparse "
    "cells require caution."
)
add_two_panel_figures(
    MORTGAGE_HEATMAP,
    "Figure 1. Mortgage default rate by LTV and DTI tier; cell labels include rates and counts.",
    LENDING_HEATMAP,
    "Figure 2. Accepted Lending Club bad-loan rate by grade and purpose.",
)

# ---------------------------------------------------------------------------
# BODY PAGE 3 — Bubble plot and PCA
# ---------------------------------------------------------------------------
page_break()
add_heading("Multivariate evidence: redundancy and dimensionality", level=2)
add_text(
    "The mortgage bubble plot confirms that loan amount closely tracks property value and "
    "that defaults overlap nondefaults, so two variables do not form a clean boundary. PCA "
    "reaches 86.3% variance with four mortgage components and 84.1% with five Lending Club "
    "components; the first two retain only 52.8% and 46.8%. Logistic ROC AUCs of 0.586 and "
    "0.680 likewise leave room for nonlinear models without guaranteeing improvement."
)
add_two_panel_figures(
    MORTGAGE_BUBBLE,
    "Figure 3. Mortgage loan amount versus property value; bubble size is DTI and color is Status.",
    PCA_FIGURE,
    "Figure 4. PCA cumulative variance; four mortgage and five Lending components exceed 80%.",
)

# ---------------------------------------------------------------------------
# BODY PAGE 4 — Required Section 2b and model table
# ---------------------------------------------------------------------------
page_break()
add_heading("b. Random forest analysis", level=2)
add_text(
    "A stratified 80/20 train-test split was held fixed for evaluation. RandomizedSearchCV "
    "tuned each RandomForestClassifier with three stratified folds and ROC AUC as the scoring "
    "rule. Class-weight balancing addressed unequal outcomes. For Lending Club, median "
    "imputation was fitted inside each cross-validation fold. For mortgage, rate_of_interest "
    "was excluded because its missingness nearly identifies Status; modeling used 124,437 "
    "complete cases on six pre-outcome numeric features. The best mortgage forest used "
    "208 trees, depth 12, sqrt feature sampling, and minimum leaf size 13. The Lending Club "
    "forest used 289 trees, depth 8, all features per split, and minimum leaf size 12."
)

add_caption("Table 1. Validated random-forest performance on held-out test sets.")
table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [1.30, 1.35, 1.15, 0.85, 0.85, 0.85]
headers = [
    "Dataset",
    "n; adverse rate",
    "CV / test AUC",
    "Accuracy",
    "Precision",
    "Recall",
]
for idx, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
    cell.width = Inches(width)
    set_cell_text(cell, header, bold=True)
    shade_cell(cell, "D9E2F3")

rows = [
    ["Mortgage", "124,437; 16.3%", "0.730 / 0.723", "0.760", "0.345", "0.525"],
    ["Lending Club", "61,100; 13.0%", "0.695 / 0.698", "0.648", "0.215", "0.643"],
]
for row_values in rows:
    row = table.add_row()
    for cell, value, width in zip(row.cells, row_values, widths):
        cell.width = Inches(width)
        set_cell_text(cell, value)
set_table_borders(table)

add_text(
    "AUC measures ranking across thresholds, whereas precision and recall describe the "
    "specific 0.50 cutoff. The mortgage confusion matrix contains 20,974 true negatives, "
    "5,056 false positives, 2,413 false negatives, and 2,667 true positives. Lending Club "
    "contains 8,621, 4,662, 712, and 1,280, respectively. The lower Lending precision shows "
    "why the threshold must be selected from review capacity and error cost—not inherited "
    "from a software default."
)

# ---------------------------------------------------------------------------
# BODY PAGE 5 — Random forest interpretation
# ---------------------------------------------------------------------------
page_break()
add_heading("Random forest interpretation: which variables the models used", level=2)
add_text(
    "Impurity-based importance indicates how often a feature improved forest splits; it is "
    "not a causal effect. Mortgage importance is led by DTI (0.317), income (0.224), and LTV "
    "(0.207). Lending Club is concentrated in interest rate (0.623); DTI, annual income, "
    "revolving utilization, installment, FICO, and total accounts each contribute 0.069 or less."
)
add_picture(IMPORTANCE_FIGURE, width=5.60)
add_caption(
    "Figure 5. Random-forest impurity importance by dataset. Values describe model use, not causation."
)
add_text(
    "Affordability dominates mortgage risk, while an interest rate already shaped by "
    "underwriting dominates Lending Club. Module C will add held-out permutation importance "
    "and time-stability checks so a single fitted forest does not determine the narrative."
)

# ---------------------------------------------------------------------------
# BODY PAGE 6 — Random forest risk segmentation
# ---------------------------------------------------------------------------
page_break()
add_heading("Random forest interpretation: operational separation", level=2)
add_text(
    "Risk deciles translate AUC into reviewer capacity. On held-out data, mortgage default "
    "rises from 6% in the lowest-risk decile to 52% in the highest, 3.2 times the complete-case "
    "average. Lending Club rises from 3% to 30%, 2.3 times average. Both sequences are ordered."
)
add_picture(DECILE_FIGURE, width=5.60)
add_caption(
    "Figure 6. Held-out adverse-outcome rate by predicted-risk decile; decile 10 is highest risk."
)
add_text(
    "The result supports prioritized human review, not automatic denial: the top decile still "
    "contains many nondefaults and the model misses defaults below it. Calibration, time "
    "stability, reason codes, and subgroup errors remain gating tests because complex models "
    "do not remove the duty to give accurate, specific adverse-action reasons (CFPB, 2022)."
)

# ---------------------------------------------------------------------------
# BODY PAGE 7 — Required Section 2c and Section 3
# ---------------------------------------------------------------------------
page_break()
add_heading("c. Relationship to analysis from Week 1-7", level=2)
add_text(
    "Week 2 documented scale, missingness, outcome definitions, and the planned model "
    "comparison. Week 4 found right-skewed amounts and class imbalance. Week 6 found weak "
    "mortgage pairwise correlations (DTI 0.078; income −0.065), while Lending Club interest "
    "rate was strongest at 0.206 and risk rose from grade A to G. Week 7 framed the 100%+ LTV "
    "risk jump."
)
add_text(
    "Weeks 8–12 add nonlinear interactions, multidimensionality, and model comparisons. The "
    "mortgage forest improves logistic AUC from 0.586 to 0.723, but the population changes: "
    "complete cases default at 16.3% versus 24.6% overall because property value and LTV are "
    "missing for 15,098 records, 99.99% with Status = 1. This is selection, not model progress."
)
add_heading("3. Data conclusions")
add_text(
    "Expected findings were the grade gradient and the importance of interest rate, income, "
    "DTI, and LTV. Unexpected findings were weak standalone mortgage credit score, nonlinear "
    "DTI, the number of PCA components needed, and near-deterministic missingness. The evidence "
    "is shareable for exploration because imbalance and limitations are explicit and metrics "
    "were reproduced. It is not production evidence: Lending Club is accepted-only, mortgage "
    "lacks time validation, outcomes may mature differently, and current data cannot establish "
    "protected-class fairness."
)

# ---------------------------------------------------------------------------
# BODY PAGE 8 — Required Section 4, proposal framing/questions
# ---------------------------------------------------------------------------
page_break()
add_heading("4. Proposal")
add_text(
    "Module C will develop an auditable early-risk ranking system for consumer lending. The "
    "primary population will be Accepted Lending Club loans because the file is large and "
    "has issue dates for out-of-time evaluation; mortgage will be a robustness benchmark, not "
    "pooled training data. The system will prioritize monitoring or human review, not issue an "
    "automatic adverse decision."
)
add_rich(
    [
        ("Research question 1. ", {"bold": True}),
        (
            "Do models using only pre-outcome, operationally available features improve "
            "out-of-time discrimination over the current logistic baseline and Week 11 "
            "random forest?",
            {},
        ),
    ]
)
add_rich(
    [
        ("Research question 2. ", {"bold": True}),
        (
            "Can predicted probabilities reliably concentrate adverse outcomes in the "
            "highest-risk deciles while remaining calibrated enough to compare risk across "
            "months?",
            {},
        ),
    ]
)
add_rich(
    [
        ("Research question 3. ", {"bold": True}),
        (
            "Are ranking, calibration, and threshold errors stable over time and across "
            "available, legally reviewable borrower or loan segments?",
            {},
        ),
    ]
)
add_text(
    "Lending Club grade and interest rate contain signal, but test AUC 0.698 leaves room to "
    "improve; the top-decile adverse rate is already 2.3 times average. Accepted-only records "
    "cannot estimate rejected-applicant outcomes, so rejected data will describe selection, "
    "not create counterfactual labels. Industry value is focused review. Public value is "
    "documented limits, errors, and explanations aligned with CFPB expectations (CFPB, 2022)."
)

# ---------------------------------------------------------------------------
# BODY PAGE 9 — Proposal process, validation, and metrics
# ---------------------------------------------------------------------------
page_break()
add_heading("Proposed modeling and validation process", level=2)
add_text(
    "First, a data dictionary will tag every candidate variable as application-time, "
    "post-origination, outcome-derived, sensitive, or unavailable at decision time. Leakage "
    "checks will compare missingness and timing with the target. Duplicates, impossible values, "
    "target maturity, and monthly shifts will be reported. Imputation, transforms, and "
    "categorical encoding will be fitted only within training folds."
)
add_text(
    "A regularized logistic regression will remain the interpretable baseline. Random forest "
    "and histogram gradient boosting will test nonlinearities. Randomized search will operate "
    "inside expanding-window validation; the latest period remains untouched for the final "
    "test. Mortgage will test broad robustness, with no averaging across different targets."
)
add_text(
    "ROC AUC and precision-recall AUC will measure ranking; log loss, Brier score, and "
    "reliability plots will measure calibration. Precision, recall, F1, confusion matrices, "
    "recall at a fixed review rate, and top-decile lift will test operating value. Bootstrap "
    "intervals and fold-to-holdout gaps will quantify uncertainty; permutation importance and "
    "reason-code review will test explanation stability."
)
add_text(
    "Where lawful and supported, audit slices will compare AUC, calibration, false-positive "
    "rate, and recall across time and available groups. Small groups will carry intervals or "
    "be suppressed. Gains that depend on leakage, unstable subgroups, or unusable reason codes "
    "will be rejected."
)

# ---------------------------------------------------------------------------
# BODY PAGE 10 — Proposal success criteria, process, teamwork
# ---------------------------------------------------------------------------
page_break()
add_heading("Success criteria and deliverables", level=2)
add_text(
    "The primary success threshold is out-of-time Lending Club ROC AUC of at least 0.73, with "
    "an improvement of at least 0.02 over the same-split logistic baseline and no more than "
    "0.03 deterioration from cross-validation to holdout. The top decile should reach 2.5 "
    "times the adverse rate, Brier score should improve 10%, and capacity-based precision and "
    "recall must beat random selection and Week 11. Recall or false-positive-rate gaps above "
    "0.05 will trigger investigation and mitigation or documented rejection."
)
add_text(
    "Deliverables are a reproducible pipeline, data dictionary, model card, validation report, "
    "threshold-capacity table, explanation review, and decision brief. Checkpoints are "
    "(1) data/leakage audit, (2) baseline/time split, (3) tuning/calibration/error analysis, "
    "and (4) robustness and recommendation: controlled pilot, revise/retest, or stop."
)
add_heading("5. Teamwork")
add_text(
    "The supplied files identify an individual submission and contain no teammate, contract, "
    "or team-communication record. No other member can be scored without fabrication."
)
team_table = doc.add_table(rows=1, cols=3)
team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
team_headers = ["Member", "Score (1–5)", "Comment"]
team_widths = [1.35, 1.10, 4.05]
for cell, header, width in zip(team_table.rows[0].cells, team_headers, team_widths):
    cell.width = Inches(width)
    set_cell_text(cell, header, bold=True)
    shade_cell(cell, "D9E2F3")
team_row = team_table.add_row()
team_values = [
    "Calvin Baker (self)",
    "5",
    (
        "Completed weekly analyses, integrated milestones, reproduced final models, and "
        "documented limitations. Individual submission; team criteria are not applicable."
    ),
]
for cell, value, width in zip(team_row.cells, team_values, team_widths):
    cell.width = Inches(width)
    set_cell_text(
        cell,
        value,
        align=WD_ALIGN_PARAGRAPH.LEFT if cell is team_row.cells[2] else WD_ALIGN_PARAGRAPH.CENTER,
    )
set_table_borders(team_table)

# ---------------------------------------------------------------------------
# APPENDIX — starts after the 10-page body
# ---------------------------------------------------------------------------
page_break()
add_heading("Appendices")
add_heading("6. Citations/bibliography")

bibliography = [
    (
        "Baker, C. (2026). DX699O2 weekly analysis notebooks, Weeks 1–12 "
        "[Unpublished course work]."
    ),
    (
        "Consumer Financial Protection Bureau. (2022, May 26). Circular 2022-03: "
        "Adverse action notification requirements in connection with credit decisions based "
        "on complex algorithms. "
    ),
    (
        "Knaflic, C. N. (2015). Storytelling with data: A data visualization guide for "
        "business professionals. Wiley."
    ),
    (
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., "
        "Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., "
        "Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, É. (2011). Scikit-learn: "
        "Machine learning in Python. Journal of Machine Learning Research, 12, 2825–2830."
    ),
    (
        "Scikit-learn developers. (2026). RandomForestClassifier, RandomizedSearchCV, and "
        "roc_auc_score [Documentation]."
    ),
    (
        "Kaggle. (n.d.). Loan default dataset [Data set]."
    ),
    (
        "Kaggle. (n.d.). All Lending Club loan data [Data set]."
    ),
    (
        "Kaggle. (n.d.). Home Credit default risk [Data set]."
    ),
]

for entry in bibliography:
    paragraph = add_text(entry)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    if entry.startswith("Consumer Financial"):
        add_hyperlink(
            paragraph,
            "CFPB circular",
            "https://www.consumerfinance.gov/compliance/circulars/"
            "circular-2022-03-adverse-action-notification-requirements-in-connection-with-"
            "credit-decisions-based-on-complex-algorithms/",
        )
    elif entry.startswith("Scikit-learn"):
        add_hyperlink(
            paragraph,
            "Random forest documentation",
            "https://scikit-learn.org/stable/modules/generated/"
            "sklearn.ensemble.RandomForestClassifier.html",
        )
    elif entry.startswith("Kaggle. (n.d.). Loan default"):
        add_hyperlink(
            paragraph,
            "Dataset page",
            "https://www.kaggle.com/datasets/yasserh/loan-default-dataset",
        )
    elif entry.startswith("Kaggle. (n.d.). All Lending"):
        add_hyperlink(
            paragraph,
            "Dataset page",
            "https://www.kaggle.com/datasets/wordsforthewise/lending-club",
        )
    elif entry.startswith("Kaggle. (n.d.). Home"):
        add_hyperlink(
            paragraph,
            "Competition page",
            "https://www.kaggle.com/competitions/home-credit-default-risk",
        )

page_break()
add_heading("7. AI Appendix")
add_heading("Prompt supplied to the AI", level=2)
add_text(
    "“Can you now take what I have from weeks 1–7 and compile it into a Word doc according "
    "to this document? Make sure that I receive the highest grade you can and hit all rubric "
    "points.”"
)
add_heading("Representative AI output", level=2)
add_text(
    "The AI mapped the paper to the rubric's required section labels, summarized the Week "
    "1–7 progression, incorporated the existing Week 8–12 multivariate and random-forest "
    "work required by the final-project guidelines, generated publication-ready PCA, feature-"
    "importance, and risk-decile figures, and produced the formatted Word draft. It also "
    "reported the independently reproduced model results: mortgage test ROC AUC 0.723 and "
    "Lending Club test ROC AUC 0.698."
)
add_heading("How the AI output was used", level=2)
add_text(
    "AI assistance was used for rubric extraction, notebook inventory, structural editing, "
    "figure layout, code-based metric reproduction, citation organization, and Word formatting. "
    "The report does not treat AI prose as evidence. Numeric claims were checked against "
    "notebook outputs and re-fitted source-data models; discrepancies and limitations were "
    "retained, including the mortgage missingness/target coupling. The author remains "
    "responsible for verifying the submission, confirming individual-versus-team status, and "
    "ensuring that the disclosed use follows course policy."
)
add_heading("Validation record", level=2)
add_bullet(
    "Mortgage source: 148,670 rows; 24.6% overall default; 124,437 complete cases; "
    "16.3% complete-case default."
)
add_bullet(
    "Mortgage missingness audit: property value and LTV are missing in 15,098 rows, and "
    "99.99% of those rows have Status = 1; rate_of_interest missingness is perfectly coupled "
    "with Status = 1 in the supplied file."
)
add_bullet(
    "Held-out refit metrics matched Week 11 after rounding: mortgage AUC 0.723, accuracy "
    "0.760, precision 0.345, recall 0.525; Lending AUC 0.698, accuracy 0.648, precision "
    "0.215, recall 0.643."
)

page_break()
add_heading("8. Weekly graphs and homework assignments")
add_text(
    "The following notebooks are companion deliverables and should be submitted with this "
    "paper. They contain the required graphs, calculations, model outputs, and written "
    "interpretations. The original notebooks remain in the project directory; the completed "
    "Week 12 submission copy preserves the original Week12.ipynb."
)

weekly_entries = [
    (
        "Week 8 — Week8.ipynb. ",
        "Heatmaps, bubble plots, PCA, and multivariate exercises; completed outputs embedded.",
    ),
    (
        "Week 9 — Week9CalvinBakerSubmission.ipynb. ",
        "Dataset heatmaps, bubbles, PCA, logistic comparisons, and story plots; completed.",
    ),
    (
        "Week 10 — Week10.ipynb. ",
        "Cross-validation, model comparison, and fixed random forests; completed outputs embedded.",
    ),
    (
        "Week 11 — Week11.ipynb. ",
        "Randomized tuning, held-out metrics, importances, and risk deciles; completed.",
    ),
    (
        "Week 12 — Week12CalvinBakerSubmission.ipynb. ",
        "Graph comparison and Chapter 8 storytelling recreation; executed, original preserved.",
    ),
]
for label, description in weekly_entries:
    add_rich([(label, {"bold": True}), (description, {})])

page_break()
add_heading("Weeks 1–7 evidence carried into the final analysis", level=2)
early_rows = [
    ["Week 1", "Data cleaning, missingness, duplicates, encoding, and outlier checks"],
    ["Week 2", "Dataset selection, scale, outcome definitions, audience, and model plan"],
    ["Week 3", "Descriptive statistics and visualization practice"],
    ["Week 4", "Univariate distributions, skew, class imbalance, and outlier interpretation"],
    ["Week 5", "Visualization and analytic exercises supporting the project workflow"],
    ["Week 6", "Bivariate correlations, grouped default rates, redundancy, and time patterns"],
    ["Week 7", "Action-oriented LTV story graph: default risk jumps at 100%+ LTV"],
]
early_table = doc.add_table(rows=1, cols=2)
early_table.alignment = WD_TABLE_ALIGNMENT.CENTER
early_widths = [0.90, 5.60]
for cell, header, width in zip(
    early_table.rows[0].cells, ["Week", "Contribution"], early_widths
):
    cell.width = Inches(width)
    set_cell_text(cell, header, bold=True)
    shade_cell(cell, "E8EEF7")
for row_values in early_rows:
    row = early_table.add_row()
    for idx, (cell, value, width) in enumerate(zip(row.cells, row_values, early_widths)):
        cell.width = Inches(width)
        set_cell_text(
            cell,
            value,
            align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT,
        )
set_table_borders(early_table)

# Update fields automatically when Word opens the document.
settings = doc.settings.element
update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
settings.append(update_fields)

doc.save(OUTPUT)
print(OUTPUT)
